"""
Convert Project_Report_Content.md to a formatted Word Document (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def convert_md_to_docx():
    doc = Document()
    
    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # --- Define/modify styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read the markdown file
    md_path = os.path.join(os.path.dirname(__file__), 'Project_Report_Content.md')
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')

        # --- Code Block Handling ---
        if line.startswith('```'):
            if in_code_block:
                # End code block - write buffered code
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)
                code_buffer = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # --- Table Handling ---
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows like |---|---|
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False
            # Don't increment, re-process this line

        # --- Horizontal Rule ---
        if line.strip() == '---':
            # Add a thin line / just skip
            i += 1
            continue
        
        # --- Empty lines ---
        if line.strip() == '':
            i += 1
            continue

        # --- Headers ---
        if line.startswith('######'):
            p = doc.add_heading(line.lstrip('#').strip(), level=5)
            i += 1
            continue
        elif line.startswith('#####'):
            p = doc.add_heading(line.lstrip('#').strip(), level=5)
            i += 1
            continue
        elif line.startswith('####'):
            p = doc.add_heading(line.lstrip('#').strip(), level=4)
            i += 1
            continue
        elif line.startswith('###'):
            p = doc.add_heading(line.lstrip('#').strip(), level=3)
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line.lstrip('#').strip(), level=2)
            i += 1
            continue
        elif line.startswith('# '):
            p = doc.add_heading(line.lstrip('#').strip(), level=1)
            i += 1
            continue

        # --- Blockquote ---
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            _add_formatted_run(p, text, italic=True)
            i += 1
            continue

        # --- Bullet points ---
        if re.match(r'^(\s*)[-*]\s', line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r'^(\s*)[-*]\s+', '', line)
            level = min(indent // 2, 3)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
            _add_formatted_run(p, text)
            i += 1
            continue

        # --- Numbered list ---
        if re.match(r'^(\s*)\d+\.\s', line):
            text = re.sub(r'^(\s*)\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_run(p, text)
            i += 1
            continue

        # --- Normal paragraph ---
        p = doc.add_paragraph()
        _add_formatted_run(p, line)
        i += 1

    # Flush pending table
    if in_table:
        _add_table(doc, table_rows)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Neuro_Trap_Project_Report.docx')
    doc.save(output_path)
    print(f"\n✅ Word document saved successfully!")
    print(f"📄 File: {output_path}")
    print(f"📎 You can now open it in Microsoft Word and print/export to PDF.")


def _add_formatted_run(paragraph, text, italic=False):
    """Parse inline markdown (bold, italic, code, links) and add formatted runs."""
    # Pattern to find **bold**, *italic*, `code`, [text](url)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\)|([^*`\[]+))'
    
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
            if italic:
                run.italic = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif match.group(5):  # [text](url)
            run = paragraph.add_run(match.group(5))
            run.font.color.rgb = RGBColor(0, 0, 200)
            run.underline = True
        elif match.group(7):  # plain text
            run = paragraph.add_run(match.group(7))
            if italic:
                run.italic = True


def _add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows or not rows[0]:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_run(p, cell_text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
    
    doc.add_paragraph()  # spacing after table


if __name__ == '__main__':
    convert_md_to_docx()
