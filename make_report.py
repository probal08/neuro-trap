import sys
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor

INPUT = r"d:\End sem project_honeypot\neuro-trap\Project_Report_Content.md"
OUTPUT = r"d:\End sem project_honeypot\neuro-trap\Neuro_Trap_Project_Report.docx"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.17)
sec.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_rich(p, text, is_italic=False):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        else:
            r = p.add_run(part)
            if is_italic:
                r.italic = True

i = 0
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i].rstrip('\r\n')

    if line.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Consolas'
            r.font.size = Pt(8.5)
            code_lines = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Skip ---
    if line.strip() == '---':
        i += 1
        continue
    # Skip empty
    if line.strip() == '':
        i += 1
        continue

    # Table detection
    if line.strip().startswith('|') and '|' in line:
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            raw = lines[i].strip().strip('|')
            cells = [c.strip() for c in raw.split('|')]
            if not all(re.match(r'^[-:]+$', c) for c in cells):
                rows.append(cells)
            i += 1
        if rows:
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = 'Light Grid Accent 1'
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < ncols:
                        tbl.cell(ri, ci).text = cell
                        if ri == 0:
                            for run in tbl.cell(ri, ci).paragraphs[0].runs:
                                run.bold = True
            doc.add_paragraph()
        continue

    # Headers
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
    elif line.startswith('#####'):
        doc.add_heading(line.lstrip('#').strip(), level=5)
    # Blockquote
    elif line.startswith('>'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        add_rich(p, line.lstrip('>').strip(), is_italic=True)
    # Numbered list
    elif re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s+', '', line)
        p = doc.add_paragraph(style='List Number')
        add_rich(p, text)
    # Bullet
    elif re.match(r'^(\s*)[-*]\s', line):
        indent = len(line) - len(line.lstrip())
        text = re.sub(r'^(\s*)[-*]\s+', '', line)
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + (indent // 2) * 0.63)
        add_rich(p, text)
    # Normal
    else:
        p = doc.add_paragraph()
        add_rich(p, line)

    i += 1

try:
    doc.save(OUTPUT)
    print(f"SUCCESS! Report saved to: {OUTPUT}")
except Exception as e:
    print(f"ERROR saving: {e}")
    sys.exit(1)
